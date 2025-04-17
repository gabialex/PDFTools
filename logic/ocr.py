
# logic/ocr.py
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
import os
from docx import Document
from docx.enum.text import WD_BREAK

def ocr_pdf(input_path: str, output_dir: str, language: str, 
           progress_callback=None, output_format: str = 'docx') -> str:
    """OCR a PDF and save the result to a new file with proper formatting."""
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input PDF {input_path} not found")

    doc = fitz.open(input_path)
    total_pages = len(doc)

    # Choose output format
    if output_format == 'docx':
        document = Document()
    elif output_format == 'rtf':
        rtf_content = [r'{\rtf1\ansi\ansicpg1252\deff0\nouicompat{\fonttbl{\f0\fnil Arial;}}',
                      r'{\colortbl ;\red0\green0\blue0;}',
                      r'\viewkind4\uc1\pard\f0\fs24']
    else:
        raise ValueError("Unsupported format. Use 'docx' or 'rtf'")

    if progress_callback:
        progress_callback(0, total_pages)

    for page_num in range(total_pages):
        page = doc.load_page(page_num)
        pix = page.get_pixmap()
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        text = pytesseract.image_to_string(img, lang=language)

        # Add text with pagination
        if output_format == 'docx':
            paragraph = document.add_paragraph(text)
            if page_num < total_pages - 1:  # No page break after last page
                paragraph.add_run().add_break(WD_BREAK.PAGE)
        elif output_format == 'rtf':
            rtf_content.append(text.replace('\n', '\\par '))
            if page_num < total_pages - 1:
                rtf_content.append(r'\page')

        if progress_callback:
            progress_callback(page_num + 1, total_pages)

    # Generate output path
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    output_filename = f"OCR_{base_name}.{output_format}"
    output_path = os.path.join(output_dir, output_filename)

    # Save the document
    if output_format == 'docx':
        document.save(output_path)
    elif output_format == 'rtf':
        rtf_content.append('}')
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(rtf_content))

    doc.close()

    if progress_callback:
        progress_callback(total_pages, total_pages)

    return output_path
